# -*- coding: utf-8 -*-
import json
import os
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULT_PATH = os.path.join(ROOT, "test_results.json")
MD_PATH = os.path.join(ROOT, "项目研究报告.md")
DOCX_PATH = os.path.join(ROOT, "项目研究报告.docx")


def decode_text(text: str) -> str:
    try:
        return text.encode("latin1").decode("utf-8")
    except Exception:
        return text


def load_results() -> dict:
    with open(RESULT_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    for case in data.get("results", []):
        case["name"] = decode_text(case["name"])
        case["query"] = decode_text(case["query"])
        case["vector"]["answer"] = decode_text(case["vector"]["answer"])
        case["graph"]["answer"] = decode_text(case["graph"]["answer"])
        case["vector"]["sources"] = [decode_text(item) for item in case["vector"].get("sources", [])]
        case["graph"]["sources"] = [decode_text(item) for item in case["graph"].get("sources", [])]

    return data


def summarize_case(case: dict) -> str:
    name = case["name"]
    if name == "简单查询":
        return "纯向量 RAG 与 GraphRAG 都能回答出《功夫》的导演是周星驰，说明在直接事实类问题上，两种方法都具备基本可用性。"
    if name == "多跳推理":
        return "纯向量 RAG 主要依赖召回片段做语义拼接，能够给出正确方向的回答；GraphRAG 在答案中引入了“朱茵-饰演-紫霞仙子-出现于《大话西游》”等图谱关系，回答更完整，更适合多跳问题。"
    return "在知识库中不存在“奥斯卡奖次数”的直接证据时，两种方法都没有强行编造答案；GraphRAG 会明确指出依据是当前图谱与文档证据不足，拒答边界更清晰。"


def build_markdown(data: dict) -> str:
    cases_md = []
    for case in data["results"]:
        graph_relations = case["graph"].get("graph", [])
        relation_preview = "\n".join(
            [f"- ({item['s']}) -[{item['r']}]-> ({item['o']})" for item in graph_relations[:8]]
        ) or "- 本案例未命中有效图谱关系。"

        cases_md.append(
            f"""### {case['name']}

问题：{case['query']}

纯向量 RAG 回答：
{case['vector']['answer']}

GraphRAG 回答：
{case['graph']['answer']}

GraphRAG 命中的图谱关系摘要：
{relation_preview}

效果分析：
{summarize_case(case)}
"""
        )

    return f"""# GraphRAG 增强型知识图谱问答系统研究报告

## 1. 项目概述

本项目面向课程中“增强型 GraphRAG 系统”的实现要求，构建了一个基于 React、Express、SQLite 与智谱大模型的知识图谱问答系统。系统支持文本知识导入、实体关系抽取、向量索引构建、图谱检索以及自然语言问答，前端提供对话交互与图谱可视化界面，后端提供数据入库、图谱构建和混合检索接口。

## 2. 系统需求描述

### 2.1 功能需求

1. 支持用户导入非结构化文本知识。
2. 支持自动抽取实体和关系，构建知识图谱。
3. 支持对知识文本生成向量表示，用于相似度检索。
4. 支持用户提出自然语言问题，并基于知识库生成答案。
5. 支持同时展示回答文本、参考证据片段与图谱结构。
6. 支持清空数据库和重新导入样本知识。
7. 支持纯向量 RAG 与 GraphRAG 两种模式的对照测试。

### 2.2 非功能需求

1. 系统应具备可视化交互界面，操作流程清晰。
2. 系统应具备一定的可扩展性，便于替换模型与增加数据源。
3. 系统应能在本地环境中直接运行，降低部署成本。

## 3. 系统设计

### 3.1 总体架构

系统采用“前端 + 后端 + 数据层 + 模型层”的四层结构：

1. 前端层：基于 React 构建对话页面与图谱可视化页面，负责用户输入、样本导入、结果展示。
2. 后端层：基于 Express 提供 REST API，承担知识入库、检索与问答调度。
3. 数据层：基于 SQLite 存储文档、实体和关系数据，并保存向量。
4. 模型层：调用智谱开放平台完成 embedding、三元组抽取、实体识别和答案生成。

### 3.2 模块划分

1. 知识入库模块：接收文本，生成向量，抽取三元组，并写入数据库。
2. 实体对齐模块：对新旧实体进行同义项映射，减少重复节点。
3. 向量检索模块：根据用户问题检索语义相近的文档片段。
4. 图谱检索模块：识别查询实体，并进行 2-Hop 图谱检索。
5. 问答生成模块：将图谱关系与文档片段整合后提交给智谱生成答案。
6. 对照测试模块：新增 `/api/query-vector` 接口，仅走 embedding 检索，不使用图谱关系。

## 4. 系统实现

### 4.1 前端实现

前端位于 `src/App.tsx`，主要实现了以下内容：

1. 对话输入与结果展示。
2. 示例知识库一键导入。
3. 知识库清空操作。
4. 图谱节点和关系的 D3 可视化展示。
5. 文档证据片段展示。

### 4.2 后端实现

后端入口为 `server.mjs`，核心接口包括：

1. `POST /api/ingest`：知识导入接口。
2. `POST /api/query`：GraphRAG 混合检索接口。
3. `POST /api/query-vector`：纯向量 RAG 对照接口。
4. `GET /api/graph`：返回当前图谱结构。
5. `POST /api/clear`：清空数据库。

### 4.3 核心流程

1. 文本导入后，系统先调用智谱 embedding 接口生成文本向量。
2. 系统调用知识抽取 Prompt 生成三元组。
3. 若模型抽取不稳定，系统会使用规则兜底补充结构化三元组，以保证样例知识可以稳定建图。
4. 查询时，系统同时进行向量召回和图谱检索。
5. GraphRAG 会把图谱关系与文档片段一起送入问答模型；纯向量 RAG 只使用文档片段。

## 5. 系统效果展示

建议在本节插入以下截图：

1. 系统首页截图：展示输入框、左侧系统功能按钮、右侧图谱视图。
2. 示例知识导入完成截图：展示系统成功导入知识后的状态。
3. 简单查询效果截图：展示问答结果与文档证据片段。
4. 多跳推理截图：展示 GraphRAG 右侧图谱关系命中情况。
5. 边界案例截图：展示系统拒绝编造答案的效果。

可在截图下方分别标注“图 1 系统首页”“图 2 知识导入完成界面”等说明文字。

## 6. 核心 Prompt 展示

### 6.1 知识抽取 Prompt

```text
你是知识图谱抽取助手。请只输出 JSON 数组，不要输出解释。
每个元素必须是 {{"s":"主语","r":"关系","o":"宾语"}}。
如果没有明确关系，返回 []。

请从下面文本中抽取实体关系三元组：
{{TEXT}}
```

### 6.2 实体对齐 Prompt

```text
你是实体对齐助手。请只输出 JSON 对象，key 是新实体名，value 是已存在实体名。
只有在两个名称明确指向同一实体时才映射，否则不要包含该 key。

新实体：{{NEW_ENTITIES}}
已存在实体：{{EXISTING_ENTITIES}}
```

### 6.3 查询实体识别 Prompt

```text
你是查询实体识别助手。请只输出 JSON 字符串数组，
提取用户问题中的关键实体名；如果没有明显实体，返回 []。

{{QUERY}}
```

### 6.4 GraphRAG 问答生成 Prompt

```text
你是 GraphRAG 问答助手。请优先基于给定的图谱关系和文档片段作答，回答使用中文；
如果证据不足，请明确说明不确定，不要编造。

知识图谱上下文：
{{GRAPH_CONTEXT}}

文档片段：
{{DOCUMENT_CONTEXT}}

用户问题：
{{QUERY}}
```

### 6.5 纯向量 RAG 问答 Prompt

```text
你是纯向量RAG问答助手。你只能基于给定文档片段作答，回答使用中文；
如果文档中没有足够证据，请明确说明无法回答，不要补充图谱推理和外部常识。

文档片段：
{{DOCUMENT_CONTEXT}}

用户问题：
{{QUERY}}
```

## 7. 测试报告

### 7.1 测试环境

1. 操作系统：Windows
2. 前端：React + Vite
3. 后端：Express + SQLite
4. 大模型与向量模型：智谱开放平台
5. 测试方式：新增 `/api/query-vector` 与 `/api/query` 两个接口进行对照测试

### 7.2 测试结果

{chr(10).join(cases_md)}

## 8. 结论

通过本次实现可以看出，纯向量 RAG 在直接事实型问题上已经能够提供较好的结果，但在多跳关系表达、结构化解释和证据链展示方面，GraphRAG 更具优势。与此同时，当前图谱构建质量仍会受到抽取与实体对齐精度影响，因此后续仍可继续优化三元组抽取质量、关系去重、实体规范化和图谱约束策略。

## 9. 复现实验方式

1. 启动系统：`npm run dev`
2. 导入示例知识库
3. 使用 `/api/query-vector` 测试纯向量 RAG
4. 使用 `/api/query` 测试 GraphRAG
5. 运行 `python scripts/run_eval.py` 可自动生成 `test_results.json`
"""


def build_docx(markdown_text: str, data: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run("GraphRAG 增强型知识图谱问答系统研究报告")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("含核心 Prompt、系统设计说明与测试对比报告")

    doc.add_paragraph("")

    for line in markdown_text.splitlines():
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("```"):
            continue
        if line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. ") or line.startswith("6. ") or line.startswith("7. "):
            doc.add_paragraph(line, style="List Number")
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
            continue
        if line.strip():
            doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("附录：测试结果对照表", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "测试类型"
    header[1].text = "测试问题"
    header[2].text = "纯向量 RAG"
    header[3].text = "GraphRAG"
    header[4].text = "结论"

    for case in data["results"]:
        row = table.add_row().cells
        row[0].text = case["name"]
        row[1].text = case["query"]
        row[2].text = case["vector"]["answer"]
        row[3].text = case["graph"]["answer"]
        row[4].text = summarize_case(case)

    doc.save(DOCX_PATH)


def main() -> None:
    data = load_results()
    markdown_text = build_markdown(data)
    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write(markdown_text)
    build_docx(markdown_text, data)
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
